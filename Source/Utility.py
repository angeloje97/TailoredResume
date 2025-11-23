from pathlib import Path
from icecream import ic
from docx import Document
from docx2pdf import convert
from pygame import mixer
import json
import shutil
import os
from datetime import datetime

#region Global Variables

base_dir = Path(__file__).parent.parent

paths = {
    "base_resume": base_dir / "BaseResumes",
    "resources": base_dir / "Resources",
    "results": base_dir / "Results",
    "json_data" : base_dir / "Resources" / "Json Data",
    "temp": base_dir / "Temp"
}

base_resumes = []
base_resume_texts = []
full_base_resume_text = ""
resume_prompt = ""


json_template = {}
resume_template = ""
cover_letter_template = ""

config = {}

#endregion

#region Functions

def get_base_resumes():
    global paths
    global base_resumes

    folder = Path(paths["base_resume"])
    ensure_path_exists(folder)

    base_resumes = [f for f in folder.glob("*.docx") if not f.name.startswith("~$")]

def scan_docx(docx_path, modifier) -> Document:
    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        modifier(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                modifier(cell.text)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                modifier(paragraph.text)

        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                modifier(paragraph.text)

    return doc

def modify_docx(docx_path, modifier) -> Document:
    doc = Document(docx_path)

    def process_paragraph(paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)

        modified_text = modifier(full_text)
        
        applied_text = False
        for run in paragraph.runs:
            run.text = ""

            if not applied_text:
                applied_text = True
                paragraph.runs[0].text = modified_text
        
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
                        
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                process_paragraph(paragraph)
                    

        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                process_paragraph(paragraph)
                    

    return doc

def get_docx_text(docx_path):

    full_text = []

    scan_docx(docx_path, lambda s: full_text.append(s))

    return full_text

def get_templates():
    global paths
    global resume_template
    global json_template
    global cover_letter_template
    global resume_prompt

    prompt_path = paths['resources'] / "Resume Prompt.md"
    json_path = paths['resources'] / "Json Template.json"
    resume_template = paths['resources'] / "Resume Template.docx"
    cover_letter_template = paths['resources'] / "Cover Letter Template.docx"

    with open(json_path, "r", encoding="utf-8") as file:
        json_template = json.load(file)
    
    with open(prompt_path, 'r', encoding="utf-8") as file:
        resume_prompt = file.read()

def get_json_datas():
    global paths
    json_datas = []
    ensure_path_exists(paths["json_data"])

    # Only get JSON files in the json_data folder, not in subdirectories (like Archived)
    json_paths = [p for p in paths['json_data'].glob("*.json") if p.is_file()]

    for path in json_paths:
        with open(path, 'r', encoding="utf-8") as file:
            json_datas.append(json.load(file))

    return json_datas

def get_archived_datas():
    global paths

    archived_path = paths['json_data'] / "Archived"

    ensure_path_exists(archived_path)

    json_paths = [p for p in archived_path.glob("*.json") if p.is_file()]
    datas = []
    for path in json_paths:
        with open(path, 'r', encoding='utf-8') as file:
            datas.append(json.load(file))
    return datas

def archive_expired_datas():
    global paths

    expired_datas = []
    current_date = datetime.now()

    ensure_path_exists(paths['json_data'])

    json_paths = [p for p in paths['json_data'].glob("*.json") if p.is_file()]

    for path in json_paths:
        with open(path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            

            try:
                expected_date = data['Job']['Expected Response Date']
                date = datetime.strptime(expected_date, "%m/%d/%y")

                if current_date > date:
                    expired_datas.append(path.stem)
            except Exception as e:
                print("Does not have expected response date")

    for file_name in expired_datas:
        archive_json_data(file_name)

def restore_archive_data(json_file_name):
    global paths

    full_path = paths['json_data'] / f'Archived' / f"{json_file_name}.json"
    back_up_path = paths['json_data'] / f'Archived' / f"{json_file_name} Data.json"

    destination_path = paths['json_data']

    try:
        shutil.move(full_path, destination_path)
        
    except Exception as e1:
        print(f"Something went wrong with moving {json_file_name}, using backup now\n{e1}")
        try:
            shutil.move(back_up_path, destination_path)

        except Exception as e2:
            print(f"Something went wrong when moving backup path\n{e2}")


def archive_json_data(json_file_name):
    full_path = paths['json_data'] / f"{json_file_name}.json"
    destination_path = paths['json_data'] / 'Archived'

    backup_path = paths['json_data'] / f"{json_file_name} Data.json"

    ensure_path_exists(destination_path)
    try:

        shutil.move(full_path, destination_path)

    except Exception as e1:

        print(f"Trouble moving{json_file_name}\n{e1}")
        print(f"Using backup path {backup_path}")

        try:

            shutil.move(backup_path, destination_path) 

        except Exception as e2:

            print(f"Trouble moving backup path\n{e2}")


def get_config():
    global base_dir
    global config
    full_path = base_dir / "Config.json"

    with open(full_path, 'r', encoding="utf-8") as file:
        config = json.load(file)
    
    return config

def update_config(config_data):
    global base_dir
    full_path = base_dir / "Config.json"

    with open(full_path, 'w', encoding='utf-8') as file:
        json.dump(config_data, file, indent=4, ensure_ascii=False)




def save_document_result(doc: Document, name: str):
    global paths

    full_path = paths["results"] / f'{name}.docx'

    doc.save(full_path)

def save_document_temp(doc: Document, name: str):
    global paths

    ensure_path_exists(paths['temp'])

    full_path = paths['temp'] / f'{name}.docx'

    doc.save(full_path)

def copy_files(path_a, path_b):
    ensure_path_exists(path_a)
    ensure_path_exists(path_b)

    for file_name in os.listdir(path_a):
        src_path = os.path.join(path_a, file_name)
        dst_path = os.path.join(path_b, file_name)

        if os.path.isfile(src_path):
            shutil.copy2(src_path, dst_path)

def copy_temp_to_results():
    global paths

    copy_files(paths['temp'], paths['results'])

def clear_temp():
    global paths
    temp_path = paths['temp']

    for file in Path(temp_path).glob("*"):
        try:
            if file.is_file():
                file.unlink()
        except Exception as e:
            print(f"Could not remove {file.name}\n{e}")

def convert_temp_to_pdf():
    global paths
    
    path = paths['temp']

    for file in os.listdir(path):
        if not file.lower().endswith(".docx"):
            continue

        input_path = os.path.join(path, file)
        output_path = os.path.join(path, file.replace(".docx", ".pdf"))
        
        convert(input_path, output_path)

def save_json_obj(obj, file_name):
    global paths

    ensure_path_exists(paths["json_data"])
    full_path = paths["json_data"] / f"{file_name}.json"


    with open(full_path, 'w', encoding='utf-8') as file:
        json.dump(obj, file, indent=4, ensure_ascii=False)

def expand_list_to_keys(obj, seperator=""):
    result = {}

    for key, value in obj.items():
        if not isinstance(value, list):
            result[key] = value
            continue

        for i, item in enumerate(value, start = 1):
            result[f"{key}{seperator}{i}"] = item

        
    return result

def ensure_path_exists(path):
    folder_path = Path(path)
    folder_path.mkdir(exist_ok=True)

def replace_keys(dict_obj, modifier, ignore_keys=[]):
    result = {}

    for key, value in dict_obj.items():
        if key in ignore_keys:
            result[key] = value
            continue
        new_key = modifier(key)
        result[new_key] = value

    return result

def write_to_docx(template_path: str, data: object) -> Document:
    modified_data = replace_keys(data, lambda s: "{" + s + "}")
    
    return modify_docx(template_path, lambda s: fill_template(s, modified_data))

def fill_template(template_string: str, data: object) -> str:
    new_string = template_string
    for key, val in data.items():
        if not isinstance(val, str):
            continue
        new_string = new_string.replace(key, val)

    return new_string

def get_resume_full_resume_text() -> str:
    global full_base_resume_text
    global base_resumes

    get_base_resumes()

    full_base_resume_text = ''

    for resume in base_resumes:
        full_base_resume_text += f"\n{'-'*50}\n{resume.name}\n {'-'*50}\n"
        text = get_docx_text(resume)
        base_resume_texts.append('\n'.join(text))
        full_base_resume_text += "\n".join(text)

    return full_base_resume_text

def play_sound(mp3_path: str):
    mixer.init()
    mixer.music.load(mp3_path)
    mixer.music.play()

def play_notification_sound():
    full_path = paths['resources'] / "Notification Sound.mp3"

    play_sound(full_path)
#endregion

#region Main Script

get_base_resumes()
get_templates()
get_resume_full_resume_text()

# Populate base resume texts

#endregion

get_config()

get_json_datas()

copy_temp_to_results()

if(config['Settings']['Auto Archive Expired Applications']):
    archive_expired_datas()